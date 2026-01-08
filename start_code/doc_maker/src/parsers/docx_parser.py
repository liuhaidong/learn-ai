from docx import Document
from typing import List, List[List[str]]


def parse_docx(path: str) -> List[str]:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return paragraphs


def parse_docx_tables(path: str) -> List[List[List[str]]]:
    doc = Document(path)
    tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        tables.append(rows)
    return tables
