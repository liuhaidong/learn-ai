from docx import Document
from typing import List, Dict


def export_word(text: str, output_path: str = "output/final_bid.docx"):
    from pathlib import Path
    
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.isdigit() and len(line) <= 2:
            pass
        else:
            doc.add_paragraph(line)
    
    doc.save(output_path)
    return output_path
